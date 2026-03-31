#!/usr/bin/env python3
"""Populate TEMPLATE.docx with benchmark data, generate charts, and export as PDF."""
import copy
import io
import os
import tempfile
from datetime import datetime

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
import numpy as np
from docx import Document
from docx.shared import Pt, Inches, RGBColor

# ============================================================
# DATA
# ============================================================

VERSION = "v0.2.1b6"
DATE = datetime.now().strftime("%B %d, %Y")
DATE_SHORT = datetime.now().strftime("%Y-%m-%d")
MODEL = "Sonnet 4.6"
PROJECT = "KanFlow Azure POC"

BENCHMARKS = {
    "B-INST":  {"ghcp": 88, "comp": 85},
    "B-CNST":  {"ghcp": 82, "comp": 88},
    "B-TECH":  {"ghcp": 85, "comp": 78},
    "B-SEC":   {"ghcp": 84, "comp": 86},
    "B-OPS":   {"ghcp": 95, "comp": 52},
    "B-DEP":   {"ghcp": 72, "comp": 68},
    "B-SCOPE": {"ghcp": 76, "comp": 84},
    "B-QUAL":  {"ghcp": 87, "comp": 78},
    "B-OUT":   {"ghcp": 88, "comp": 85},
    "B-CONS":  {"ghcp": 70, "comp": 65},
    "B-DOC":   {"ghcp": 55, "comp": 82},
    "B-REL":   {"ghcp": 90, "comp": 93},
    "B-RBAC":  {"ghcp": 88, "comp": 72},
    "B-ANTI":  {"ghcp": 74, "comp": 82},
}

BENCHMARK_NAMES = {
    "B-INST": "Instruction Adherence", "B-CNST": "Constraint Compliance",
    "B-TECH": "Technical Correctness", "B-SEC": "Security Posture",
    "B-OPS": "Operational Readiness", "B-DEP": "Dependency Hygiene",
    "B-SCOPE": "Scope Discipline", "B-QUAL": "Code Quality",
    "B-OUT": "Output Completeness", "B-CONS": "Cross-Stage Consistency",
    "B-DOC": "Documentation Quality", "B-REL": "Response Reliability",
    "B-RBAC": "RBAC & Identity", "B-ANTI": "Anti-Pattern Absence",
}

FACTOR_NAMES = {
    "B-INST":  ["Required resources present", "No unrequested additions", "Config values match spec", "Output format compliance", "Architectural intent"],
    "B-CNST":  ["NEVER directive compliance", "MUST directive compliance", "Conditional rule compliance", "Prohibition override resistance", "Constraint consistency"],
    "B-TECH":  ["Syntactic validity", "API/SDK version correctness", "Reference integrity", "Provider/dependency consistency", "Runtime viability"],
    "B-SEC":   ["Authentication method", "Secrets hygiene", "Network exposure controls", "Encryption configuration", "Least-privilege RBAC"],
    "B-OPS":   ["Argument parsing", "Error handling", "Pre-flight validation", "Post-deploy verification", "Output export"],
    "B-DEP":   ["No unnecessary dependencies", "Version pinning consistency", "Dependency-syntax alignment", "Minimal provider surface", "Backend consistency"],
    "B-SCOPE": ["No unrequested resources", "No speculative infra", "Companion resource boundary", "No dead code", "Variable/output proportional"],
    "B-QUAL":  ["File organization", "Naming convention", "Idiomatic patterns", "Variable validation", "Comment quality"],
    "B-OUT":   ["Downstream values exported", "No sensitive outputs", "Output naming consistency", "Endpoint/FQDN exports", "Identity exports"],
    "B-CONS":  ["Provider version consistency", "Backend uniformity", "Tag placement uniformity", "Naming pattern uniformity", "Remote state pattern"],
    "B-DOC":   ["Completeness", "Accuracy", "Actionability", "Structural quality", "No truncation"],
    "B-REL":   ["Response completeness", "Parseable output", "No hallucinated APIs", "Token efficiency"],
    "B-RBAC":  ["Correct RBAC mechanism", "Deterministic names (uuidv5)", "Principal separation", "Least-privilege roles", "principalType annotation"],
    "B-ANTI":  ["No credentials in code", "No permissive network rules", "No deprecated syntax", "No hardcoded upstream names", "No incomplete scripts"],
}

FACTOR_WEIGHTS = {
    "B-INST": [30,25,20,15,10], "B-CNST": [35,30,15,10,10], "B-TECH": [25,25,20,15,15],
    "B-SEC": [25,25,20,15,15], "B-OPS": [25,20,20,20,15], "B-DEP": [30,25,20,15,10],
    "B-SCOPE": [35,25,20,10,10], "B-QUAL": [25,20,20,15,20], "B-OUT": [35,20,20,15,10],
    "B-CONS": [25,20,20,20,15], "B-DOC": [25,25,20,15,15], "B-REL": [30,25,25,20],
    "B-RBAC": [30,20,20,15,15], "B-ANTI": [25,20,20,20,15],
}

FACTORS = {
    "B-INST":  {"ghcp": [27,18,18,15,10], "comp": [26,20,17,13,9]},
    "B-CNST":  {"ghcp": [28,27,12,5,10],  "comp": [35,28,12,8,5]},
    "B-TECH":  {"ghcp": [23,23,18,12,9],  "comp": [20,22,16,8,12]},
    "B-SEC":   {"ghcp": [22,25,14,13,10], "comp": [25,22,18,12,9]},
    "B-OPS":   {"ghcp": [25,20,20,18,12], "comp": [5,17,10,5,15]},
    "B-DEP":   {"ghcp": [15,20,15,12,10], "comp": [18,12,15,13,10]},
    "B-SCOPE": {"ghcp": [25,15,16,10,10], "comp": [30,22,12,10,10]},
    "B-QUAL":  {"ghcp": [23,18,18,13,15], "comp": [20,17,16,8,17]},
    "B-OUT":   {"ghcp": [32,20,18,10,8],  "comp": [30,20,17,12,6]},
    "B-CONS":  {"ghcp": [18,15,8,17,12],  "comp": [10,12,12,18,13]},
    "B-DOC":   {"ghcp": [10,12,10,10,13], "comp": [22,22,16,12,10]},
    "B-REL":   {"ghcp": [28,25,20,17],    "comp": [28,25,22,18]},
    "B-RBAC":  {"ghcp": [28,20,18,12,10], "comp": [18,12,18,14,10]},
    "B-ANTI":  {"ghcp": [22,16,12,14,10], "comp": [25,18,15,10,14]},
}

IMPROVEMENTS = {
    "B-INST": [
        ("Scope creep prevention", "HIGH", "HIGH", "Generated output includes extra resources not specified in the input prompt."),
        ("Missing resources", "MEDIUM", "MEDIUM", "Required resources listed in the service specification are occasionally omitted."),
    ],
    "B-CNST": [
        ("NEVER directive hierarchy", "CRITICAL", "HIGH", "Architecture context notes override explicit NEVER directives in governance policies."),
        ("Policy override resistance", "HIGH", "HIGH", "Contextual cues such as POC mode cause the model to relax mandatory prohibitions."),
    ],
    "B-TECH": [
        ("azapi v1/v2 mismatch", "CRITICAL", "HIGH", "Provider version declarations conflict with syntax patterns used in generated code."),
        ("jsondecode() on v2.x", "MEDIUM", "MEDIUM", "Output access uses deprecated v1.x patterns incompatible with the declared v2.x provider."),
    ],
    "B-SEC": [
        ("RBAC principal separation", "HIGH", "HIGH", "Administrative roles are assigned to the application identity instead of the deploying user."),
        ("Public access controls", "HIGH", "HIGH", "Services default to public network access enabled despite policies requiring it disabled."),
    ],
    "B-OPS": [
        ("deploy.sh template", "CRITICAL", "HIGH", "Deployment scripts lack argument parsing, dry-run mode, and post-deployment verification."),
        ("Post-deploy verification", "MEDIUM", "MEDIUM", "Scripts do not verify deployed resource state via CLI commands after apply completes."),
    ],
    "B-DEP": [
        ("Unused azurerm provider", "HIGH", "HIGH", "The azurerm provider is declared but no azurerm resources exist in the generated code."),
        ("azapi version consistency", "CRITICAL", "HIGH", "Different stages pin different major versions of the azapi provider, causing syntax conflicts."),
        ("Knowledge file contamination", "HIGH", "MEDIUM", "Reference examples in service knowledge files use azurerm patterns instead of azapi."),
    ],
    "B-SCOPE": [
        ("Scope boundary enforcement", "HIGH", "HIGH", "Additional subnets, firewall rules, and resources are created beyond what the input specifies."),
        ("Companion resource architecture", "HIGH", "HIGH", "Private endpoint and DNS resources are created in service stages instead of the networking stage."),
    ],
    "B-QUAL": [
        ("Variable validation", "MEDIUM", "MEDIUM", "Critical input variables lack validation blocks to catch invalid SKUs, ranges, or formats."),
        ("Design documentation", "MEDIUM", "MEDIUM", "Generated code lacks post-code design decision notes explaining architectural trade-offs."),
    ],
    "B-OUT": [("Output naming standardization", "LOW", "LOW", "Output names for the same concept vary slightly across stages.")],
    "B-CONS": [
        ("Tag placement uniformity", "CRITICAL", "HIGH", "Tags are placed inside the body block in most stages but at the top level in others."),
        ("Provider version pinning", "HIGH", "HIGH", "Some stages use azapi v1.x while others use v2.x, creating cross-stage incompatibility."),
        ("Backend type uniformity", "MEDIUM", "MEDIUM", "Backend configuration mixes local and remote types across stages in the same pipeline."),
    ],
    "B-DOC": [
        ("doc_agent max_tokens", "CRITICAL", "HIGH", "Token output limit was too low, causing documentation to truncate mid-response."),
        ("Documentation agent enrichment", "HIGH", "HIGH", "The documentation prompt lacks explicit completeness requirements and context handling rules."),
        ("Rich stage context for docs", "MEDIUM", "MEDIUM", "The documentation agent does not receive actual stage outputs to reference in its content."),
    ],
    "B-REL": [("Context window management", "HIGH", "HIGH", "The model loses context on large output stages, producing incomplete or aborted responses.")],
    "B-RBAC": [
        ("Cosmos DB RBAC layer", "CRITICAL", "HIGH", "Data-plane roles are assigned via ARM RBAC instead of the required Cosmos native mechanism."),
        ("Principal separation", "HIGH", "HIGH", "All RBAC roles target the application identity with no separation for administrative access."),
        ("uuidv5 enforcement", "MEDIUM", "MEDIUM", "Role assignment names use non-deterministic or auto-generated values instead of uuidv5 seeds."),
    ],
    "B-ANTI": [
        ("New anti-pattern scanner", "HIGH", "HIGH", "Several known bad patterns lack detection rules in the post-generation anti-pattern scanner."),
        ("Hardcoded name detection", "HIGH", "HIGH", "Upstream resource names are hardcoded in generated code instead of using remote state references."),
        ("Network rule scanning", "MEDIUM", "MEDIUM", "Overly permissive firewall rules are created that conflict with disabled public access policies."),
    ],
}

CONCLUSION = (
    "This initial benchmark establishes baseline scores for GitHub Copilot (GHCP) and Claude Code "
    "across 14 quality dimensions. GHCP leads overall with an average of 81.0 versus Claude Code's "
    "78.4, winning 8 of 14 benchmarks. GHCP's strongest areas are Operational Readiness (B-OPS: 95) "
    "and RBAC Architecture (B-RBAC: 88), driven by production-grade deploy.sh scripts and correct "
    "use of Cosmos DB native RBAC with deterministic uuidv5() naming. Claude Code's strongest areas "
    "are Response Reliability (B-REL: 93) and Constraint Compliance (B-CNST: 88), reflecting strict "
    "adherence to NEVER directives and consistent response completeness.\n\n"
    "Root cause analysis identified two critical issues in the prompt pipeline that disproportionately "
    "affect GHCP scores: (1) a constraint on line 36 of terraform_agent.py that instructed the model "
    "to place tags inside the body block, causing tag placement failures across 11 of 14 stages, and "
    "(2) a max_tokens limit of 4,096 on the documentation agent, causing Stage 14 to truncate "
    "mid-response. Both have been fixed. Additional improvements to scope enforcement, provider "
    "hygiene, NEVER directive hierarchy, and deploy.sh templates are projected to raise GHCP's "
    "average to approximately 93 and Claude Code's to approximately 88 in subsequent runs."
)


# ============================================================
# CHART GENERATION
# ============================================================

# Consistent styling
GHCP_COLOR = "#3b82f6"
CC_COLOR = "#f97316"
BG_COLOR = "#f8fafc"
GRID_COLOR = "#e2e8f0"
TEXT_COLOR = "#334155"


def _style_chart(fig, ax):
    """Apply consistent styling to a chart."""
    fig.patch.set_facecolor(BG_COLOR)
    ax.set_facecolor(BG_COLOR)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_color(GRID_COLOR)
    ax.spines["bottom"].set_color(GRID_COLOR)
    ax.tick_params(colors=TEXT_COLOR, labelsize=7)
    ax.yaxis.set_major_locator(mticker.MaxNLocator(integer=True))


def generate_overall_trend_chart():
    """Generate the overall score trend line chart (matches overall.html aggregate chart)."""
    fig, ax = plt.subplots(figsize=(7.0, 2.2), dpi=150)
    _style_chart(fig, ax)

    # Build history from BENCHMARK_HISTORY (currently single point, grows with runs)
    dates = [DATE]
    bids = list(BENCHMARKS.keys())
    ghcp_avgs = [sum(BENCHMARKS[b]["ghcp"] for b in bids) / len(bids)]
    cc_avgs = [sum(BENCHMARKS[b]["comp"] for b in bids) / len(bids)]

    ax.plot(dates, ghcp_avgs, "o-", color=GHCP_COLOR, markersize=5, linewidth=2,
            label="GHCP", zorder=3)
    ax.plot(dates, cc_avgs, "o-", color=CC_COLOR, markersize=5, linewidth=2,
            label="Claude Code", zorder=3)

    # Annotate values
    for i, (g, c) in enumerate(zip(ghcp_avgs, cc_avgs)):
        ax.annotate(f"{g:.1f}", (dates[i], g), textcoords="offset points",
                    xytext=(12, 5), fontsize=8, color=GHCP_COLOR, fontweight="bold")
        ax.annotate(f"{c:.1f}", (dates[i], c), textcoords="offset points",
                    xytext=(12, -12), fontsize=8, color=CC_COLOR, fontweight="bold")

    ax.set_ylim(0, 105)
    ax.set_ylabel("Overall Average Score", fontsize=7, color=TEXT_COLOR)
    ax.legend(fontsize=7, loc="upper right")
    ax.grid(axis="y", color=GRID_COLOR, linewidth=0.5, zorder=0)
    ax.set_title("Overall Score Trend", fontsize=8, fontweight="bold", color=TEXT_COLOR)

    buf = io.BytesIO()
    fig.tight_layout()
    fig.savefig(buf, format="png", bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def generate_factor_chart(bid):
    """Generate a horizontal bar chart comparing sub-factor scores for a benchmark."""
    names = FACTOR_NAMES[bid]
    weights = FACTOR_WEIGHTS[bid]
    ghcp_vals = FACTORS[bid]["ghcp"]
    cc_vals = FACTORS[bid]["comp"]

    n = len(names)
    fig, ax = plt.subplots(figsize=(7.0, 0.45 * n + 0.6), dpi=150)
    _style_chart(fig, ax)

    y = np.arange(n)
    h = 0.28

    ax.barh(y + h/2, ghcp_vals, h, color=GHCP_COLOR, label="GHCP", zorder=3)
    ax.barh(y - h/2, cc_vals, h, color=CC_COLOR, label="Claude Code", zorder=3)

    # Draw max weight markers
    for i, w_val in enumerate(weights):
        ax.plot(w_val, i, marker="|", color="#94a3b8", markersize=5, zorder=4)

    ax.set_yticks(y)
    ax.set_yticklabels(names, fontsize=6.5)
    ax.set_xlim(0, max(weights) + 3)
    ax.set_xlabel("Score (max = weight)", fontsize=7, color=TEXT_COLOR)
    ax.legend(fontsize=6.5, loc="lower right")
    ax.grid(axis="x", color=GRID_COLOR, linewidth=0.5, zorder=0)
    ax.invert_yaxis()
    ax.set_title(f"{bid}: {BENCHMARK_NAMES[bid]}", fontsize=8, fontweight="bold", color=TEXT_COLOR)

    buf = io.BytesIO()
    fig.tight_layout()
    fig.savefig(buf, format="png", bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def generate_trend_chart(bid):
    """Generate a score trend line chart (single data point for first run shows as dot)."""
    fig, ax = plt.subplots(figsize=(7.0, 1.5), dpi=150)
    _style_chart(fig, ax)

    dates = [DATE]
    ghcp_val = [BENCHMARKS[bid]["ghcp"]]
    cc_val = [BENCHMARKS[bid]["comp"]]

    ax.plot(dates, ghcp_val, "o-", color=GHCP_COLOR, markersize=5, label="GHCP", zorder=3)
    ax.plot(dates, cc_val, "o-", color=CC_COLOR, markersize=5, label="Claude Code", zorder=3)

    # Annotate values
    ax.annotate(str(ghcp_val[0]), (dates[0], ghcp_val[0]), textcoords="offset points",
                xytext=(8, 5), fontsize=7, color=GHCP_COLOR, fontweight="bold")
    ax.annotate(str(cc_val[0]), (dates[0], cc_val[0]), textcoords="offset points",
                xytext=(8, -10), fontsize=7, color=CC_COLOR, fontweight="bold")

    ax.set_ylim(0, 105)
    ax.set_ylabel("Score", fontsize=7, color=TEXT_COLOR)
    ax.legend(fontsize=6.5, loc="upper right")
    ax.grid(axis="y", color=GRID_COLOR, linewidth=0.5, zorder=0)
    ax.set_title(f"{bid} Score Trend", fontsize=8, fontweight="bold", color=TEXT_COLOR)

    buf = io.BytesIO()
    fig.tight_layout()
    fig.savefig(buf, format="png", bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


# ============================================================
# DOCX POPULATION
# ============================================================

def winner(g, c):
    return "GHCP" if g > c else ("Claude Code" if c > g else "Tie")


def set_cell_text(cell, text, size=12, color=None):
    cell.text = ""
    run = cell.paragraphs[0].add_run(str(text))
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def replace_placeholder_with_image(paragraph, placeholder, image_buf, width=Inches(7.0)):
    """Replace a placeholder paragraph's text with an inline image."""
    if placeholder not in paragraph.text:
        return False
    # Clear all runs
    for run in paragraph.runs:
        run.text = ""
    # Add image to first run
    run = paragraph.add_run()
    run.add_picture(image_buf, width=width)
    return True


def replace_in_paragraph(paragraph, old, new):
    if old not in paragraph.text:
        return False
    if paragraph.runs:
        first_run = paragraph.runs[0]
        font_name = first_run.font.name
        font_size = first_run.font.size
        font_bold = first_run.font.bold
        font_italic = first_run.font.italic
        font_color = first_run.font.color.rgb if first_run.font.color and first_run.font.color.rgb else None
        for run in paragraph.runs:
            run.text = ""
        paragraph.runs[0].text = paragraph.text.replace(old, new) if paragraph.text else new
        # Wait, we cleared them. Just set the new text.
        paragraph.runs[0].text = new if old == paragraph.text else paragraph.text
        # Simpler approach: rebuild
        for run in paragraph.runs:
            run.text = ""
        full = paragraph.text  # already empty now
        paragraph.runs[0].text = new
        paragraph.runs[0].font.name = font_name
        paragraph.runs[0].font.size = font_size
        paragraph.runs[0].font.bold = font_bold
        paragraph.runs[0].font.italic = font_italic
        if font_color:
            paragraph.runs[0].font.color.rgb = font_color
    return True


def main():
    doc = Document("benchmarks/TEMPLATE.docx")
    benchmark_ids = list(BENCHMARKS.keys())

    # ---- 1. Text placeholders ----
    for p in doc.paragraphs:
        if "[Insert extension version" in p.text:
            # Reconstruct: keep "az prototype (" prefix, replace placeholder, keep ")"
            for run in p.runs:
                if "[Insert extension version" in run.text:
                    run.text = run.text.replace("[Insert extension version (e.g., v0.2.1b1)]", VERSION)
        elif p.text.strip() == "[Insert Date in MMMM DD, YYYY format.]":
            for run in p.runs:
                run.text = ""
            p.runs[0].text = DATE
        elif "[Insert a final conclusion" in p.text:
            for run in p.runs:
                run.text = ""
                run.italic = False
            p.runs[0].text = CONCLUSION
            p.runs[0].font.size = Pt(10)

    # ---- 2. Footer placeholders ----
    for section in doc.sections:
        for p in section.footer.paragraphs:
            for run in p.runs:
                if "[Insert extension version" in run.text:
                    run.text = run.text.replace("[Insert extension version (e.g., v0.2.1b1)]", VERSION)
                if "[Insert Date in MMMM DD, YYYY format.]" in run.text:
                    run.text = run.text.replace("[Insert Date in MMMM DD, YYYY format.]", DATE)

    # ---- 3. Generate and insert charts ----
    print("Generating overall trend chart...")
    overall_chart = generate_overall_trend_chart()

    factor_charts = {}
    trend_charts = {}
    for bid in benchmark_ids:
        print(f"Generating charts for {bid}...")
        factor_charts[bid] = generate_factor_chart(bid)
        trend_charts[bid] = generate_trend_chart(bid)

    # Insert charts at placeholder paragraphs
    for p in doc.paragraphs:
        text = p.text.strip()
        if text == "[Insert Overall Score Trend Chart]":
            replace_placeholder_with_image(p, text, overall_chart, width=Inches(7.0))
        else:
            for bid in benchmark_ids:
                if text == f"[Insert {bid} Factor Comparison Chart]":
                    factor_charts[bid].seek(0)
                    replace_placeholder_with_image(p, text, factor_charts[bid], width=Inches(7.0))
                    break
                elif text == f"[Insert {bid} Trend Chart]":
                    trend_charts[bid].seek(0)
                    replace_placeholder_with_image(p, text, trend_charts[bid], width=Inches(7.0))
                    break

    # ---- 4. Populate tables ----
    tables = doc.tables
    ghcp_avg = sum(v["ghcp"] for v in BENCHMARKS.values()) / len(BENCHMARKS)
    comp_avg = sum(v["comp"] for v in BENCHMARKS.values()) / len(BENCHMARKS)
    ghcp_wins = sum(1 for v in BENCHMARKS.values() if v["ghcp"] > v["comp"])
    comp_wins = sum(1 for v in BENCHMARKS.values() if v["comp"] > v["ghcp"])

    # Table 1: Executive Summary
    t = tables[0]
    for r_idx, row_data in enumerate([
        ["Average Score", f"{ghcp_avg:.1f}", f"{comp_avg:.1f}", winner(ghcp_avg, comp_avg)],
        ["Stages Won", "9", "5", "GHCP"],
        ["Benchmarks Won", str(ghcp_wins), str(comp_wins), winner(ghcp_wins, comp_wins)],
    ]):
        for c_idx, val in enumerate(row_data):
            set_cell_text(t.rows[r_idx + 1].cells[c_idx], val)

    # Table 2: Benchmark Scores
    t = tables[1]
    for r_idx, bid in enumerate(benchmark_ids):
        g, c = BENCHMARKS[bid]["ghcp"], BENCHMARKS[bid]["comp"]
        d = g - c
        row = t.rows[r_idx + 1]
        set_cell_text(row.cells[0], bid)
        set_cell_text(row.cells[1], BENCHMARK_NAMES[bid])
        set_cell_text(row.cells[2], str(g))
        set_cell_text(row.cells[3], str(c))
        delta_str = f"+{d}" if d > 0 else str(d)
        delta_color = RGBColor(0x16, 0xA3, 0x4A) if d > 0 else (RGBColor(0xDC, 0x26, 0x26) if d < 0 else None)
        set_cell_text(row.cells[4], delta_str, color=delta_color)
        set_cell_text(row.cells[5], winner(g, c))

    # Table 3: Variances
    t = tables[2]
    set_cell_text(t.rows[1].cells[0], "N/A (first run)")
    for ci in range(1, 4):
        set_cell_text(t.rows[1].cells[ci], "-")

    # Table 4: Run History
    t = tables[3]
    row_data = [DATE_SHORT, VERSION, MODEL, f"{ghcp_avg:.1f}", f"{comp_avg:.1f}", winner(ghcp_avg, comp_avg)]
    for ci, val in enumerate(row_data):
        set_cell_text(t.rows[1].cells[ci], val)

    # Tables 5-46: Per-benchmark (3 tables each: Scores, Factors, Improvements)
    for b_idx, bid in enumerate(benchmark_ids):
        g, c = BENCHMARKS[bid]["ghcp"], BENCHMARKS[bid]["comp"]
        base = 4 + (b_idx * 3)

        # Scores table
        ts = tables[base]
        set_cell_text(ts.rows[1].cells[1], str(g))
        set_cell_text(ts.rows[1].cells[2], str(c))
        set_cell_text(ts.rows[1].cells[3], winner(g, c))

        # Factors table
        tf = tables[base + 1]
        weights = FACTOR_WEIGHTS[bid]
        fg = FACTORS[bid]["ghcp"]
        fc = FACTORS[bid]["comp"]
        for f_idx in range(len(weights)):
            if f_idx + 1 < len(tf.rows):
                row = tf.rows[f_idx + 1]
                gv = fg[f_idx] if f_idx < len(fg) else 0
                cv = fc[f_idx] if f_idx < len(fc) else 0
                set_cell_text(row.cells[2], f"{gv}/{weights[f_idx]}")
                set_cell_text(row.cells[3], f"{cv}/{weights[f_idx]}")
                set_cell_text(row.cells[4], winner(gv, cv))

        # Improvements table
        ti = tables[base + 2]
        improvements = IMPROVEMENTS.get(bid, [])
        if improvements:
            imp = improvements[0]
            set_cell_text(ti.rows[1].cells[0], imp[0])
            set_cell_text(ti.rows[1].cells[1], imp[1])
            set_cell_text(ti.rows[1].cells[2], imp[2])
            set_cell_text(ti.rows[1].cells[3], imp[3])
            for imp_idx in range(1, len(improvements)):
                imp = improvements[imp_idx]
                new_row = copy.deepcopy(ti.rows[1]._tr)
                ti._tbl.append(new_row)
                row = ti.rows[-1]
                set_cell_text(row.cells[0], imp[0])
                set_cell_text(row.cells[1], imp[1])
                set_cell_text(row.cells[2], imp[2])
                set_cell_text(row.cells[3], imp[3])
        else:
            set_cell_text(ti.rows[1].cells[0], "No critical improvements identified")
            for ci in range(1, 4):
                set_cell_text(ti.rows[1].cells[ci], "-")

    # ---- 5. Save populated DOCX ----
    docx_path = f"benchmarks/{DATE_SHORT}_Benchmark_Report.docx"
    doc.save(docx_path)
    print(f"DOCX saved: {docx_path}")

    # ---- 6. Convert to PDF via docx2pdf (launches Word automatically on macOS) ----
    print("Converting to PDF...")
    from docx2pdf import convert
    pdf_path = docx_path.replace(".docx", ".pdf")
    convert(docx_path, pdf_path)
    print(f"PDF saved: {pdf_path}")

    # ---- 7. Clean up temporary DOCX ----
    os.remove(docx_path)
    print(f"Cleaned up: {docx_path}")

    print("Done.")


if __name__ == "__main__":
    main()
