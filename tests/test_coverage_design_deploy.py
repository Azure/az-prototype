"""Targeted tests to improve coverage for design_stage.py and deploy_stage.py.

Covers uncovered lines identified by coverage analysis:
  - design_stage.py: 102-109, 144, 318, 355-368, 389-407, 411-415, 434-438
  - deploy_stage.py: 313-321, 374, 379, 388-389, 502, 525, 669
"""

import json
from unittest.mock import MagicMock, patch

import pytest
from knack.util import CLIError

from azext_prototype.stages.discovery import DiscoveryResult
from tests.conftest import make_ai_response


# ======================================================================
# Helpers
# ======================================================================

def _make_discovery_result(**overrides):
    """Quick factory for DiscoveryResult."""
    defaults = dict(
        requirements="Build a web app",
        conversation=[],
        policy_overrides=[],
        exchange_count=2,
        cancelled=False,
    )
    defaults.update(overrides)
    return DiscoveryResult(**defaults)


def _make_agent_context(project_dir, ai_provider=None, config=None):
    from azext_prototype.agents.base import AgentContext

    return AgentContext(
        project_config=config or {"project": {"name": "test", "location": "eastus", "iac_tool": "terraform"}},
        project_dir=str(project_dir),
        ai_provider=ai_provider or MagicMock(),
    )


# ======================================================================
# DesignStage — targeted coverage
# ======================================================================


class TestDesignStageArtifactsPath:
    """Cover lines 102-109 — artifacts_path handling in execute()."""

    @patch("azext_prototype.stages.design_stage.DiscoverySession")
    def test_execute_with_artifacts_file(
        self, MockDS, project_with_config, mock_agent_context, populated_registry
    ):
        """When artifacts= points to a file, the content is ingested."""
        from azext_prototype.stages.design_stage import DesignStage

        stage = DesignStage()
        stage.get_guards = lambda: []

        artifact_file = project_with_config / "requirements.md"
        artifact_file.write_text("# Requirements\n- Feature A\n- Feature B", encoding="utf-8")

        mock_agent_context.project_dir = str(project_with_config)
        mock_agent_context.ai_provider.chat.return_value = make_ai_response("## Architecture\nDesign")

        MockDS.return_value.run.return_value = _make_discovery_result()

        prints = []
        result = stage.execute(
            mock_agent_context,
            populated_registry,
            artifacts=str(artifact_file),
            context="",
            interactive=False,
            print_fn=prints.append,
        )

        assert result["status"] == "success"
        # Verify artifacts were passed to the discovery session
        call_kwargs = MockDS.return_value.run.call_args
        assert call_kwargs is not None

    @patch("azext_prototype.stages.design_stage.DiscoverySession")
    def test_execute_with_artifacts_directory(
        self, MockDS, project_with_config, mock_agent_context, populated_registry
    ):
        """When artifacts= points to a directory, all supported files are read."""
        from azext_prototype.stages.design_stage import DesignStage

        stage = DesignStage()
        stage.get_guards = lambda: []

        art_dir = project_with_config / "specs"
        art_dir.mkdir()
        (art_dir / "overview.md").write_text("# Overview", encoding="utf-8")
        (art_dir / "data.json").write_text('{"key": "value"}', encoding="utf-8")
        (art_dir / "image.png").write_bytes(b"\x89PNG")  # unsupported extension

        mock_agent_context.project_dir = str(project_with_config)
        mock_agent_context.ai_provider.chat.return_value = make_ai_response("## Architecture")

        MockDS.return_value.run.return_value = _make_discovery_result()

        result = stage.execute(
            mock_agent_context,
            populated_registry,
            artifacts=str(art_dir),
            interactive=False,
            print_fn=lambda _: None,
        )

        assert result["status"] == "success"


class TestDesignStageNoArchitect:
    """Cover line 144 — CLIError when no architect agents found."""

    @patch("azext_prototype.stages.design_stage.DiscoverySession")
    def test_no_architect_agents_raises(self, MockDS, project_with_config, mock_agent_context):
        from azext_prototype.stages.design_stage import DesignStage
        from azext_prototype.agents.registry import AgentRegistry

        stage = DesignStage()
        stage.get_guards = lambda: []

        mock_agent_context.project_dir = str(project_with_config)
        MockDS.return_value.run.return_value = _make_discovery_result()

        empty_registry = AgentRegistry()

        with pytest.raises(CLIError, match="No architect agents available"):
            stage.execute(
                mock_agent_context,
                empty_registry,
                context="Build something",
                interactive=False,
                print_fn=lambda _: None,
            )


class TestDesignStagePolicyOverrides:
    """Cover the policy_overrides persistence path."""

    @patch("azext_prototype.stages.design_stage.DiscoverySession")
    def test_policy_overrides_stored(
        self, MockDS, project_with_config, mock_agent_context, populated_registry
    ):
        from azext_prototype.stages.design_stage import DesignStage

        stage = DesignStage()
        stage.get_guards = lambda: []

        mock_agent_context.project_dir = str(project_with_config)
        mock_agent_context.ai_provider.chat.return_value = make_ai_response("## Arch")

        MockDS.return_value.run.return_value = _make_discovery_result(
            policy_overrides=[{"policy": "managed-identity", "action": "warn"}],
        )

        result = stage.execute(
            mock_agent_context,
            populated_registry,
            interactive=False,
            print_fn=lambda _: None,
        )

        assert result["status"] == "success"
        state_path = project_with_config / ".prototype" / "state" / "design.json"
        state = json.loads(state_path.read_text(encoding="utf-8"))
        assert len(state["policy_overrides"]) == 1


class TestDesignStageIaCReview:
    """Cover line 318 — _run_iac_review method."""

    def test_run_iac_review_with_terraform_agent(self, project_with_config, mock_agent_context, populated_registry):
        from azext_prototype.stages.design_stage import DesignStage
        from azext_prototype.config import ProjectConfig

        stage = DesignStage()
        config = ProjectConfig(str(project_with_config))
        config.load()

        mock_agent_context.project_dir = str(project_with_config)

        stage._run_iac_review(
            mock_agent_context,
            populated_registry,
            config,
            MagicMock(name="cloud-architect"),
            "## Architecture Design\nUse App Service and CosmosDB",
        )

    def test_run_iac_review_no_iac_agents(self, project_with_config, mock_agent_context):
        from azext_prototype.stages.design_stage import DesignStage
        from azext_prototype.agents.registry import AgentRegistry
        from azext_prototype.config import ProjectConfig

        stage = DesignStage()
        config = ProjectConfig(str(project_with_config))
        config.load()

        empty_registry = AgentRegistry()
        # Should return early, no error
        stage._run_iac_review(
            mock_agent_context,
            empty_registry,
            config,
            MagicMock(name="cloud-architect"),
            "## Architecture",
        )


class TestDesignStageReadArtifacts:
    """Cover _read_artifacts — reads ALL files, no extension filter."""

    def test_read_artifacts_single_file(self, tmp_path):
        from azext_prototype.stages.design_stage import DesignStage

        stage = DesignStage()
        f = tmp_path / "spec.md"
        f.write_text("# Spec\nDetails here", encoding="utf-8")

        result = stage._read_artifacts(str(f))
        assert "Spec" in result["content"]
        assert len(result["read"]) == 1
        assert result["failed"] == []

    def test_read_artifacts_directory(self, tmp_path):
        from azext_prototype.stages.design_stage import DesignStage

        stage = DesignStage()
        (tmp_path / "a.md").write_text("Alpha", encoding="utf-8")
        (tmp_path / "b.yaml").write_text("key: value", encoding="utf-8")
        (tmp_path / "c.txt").write_text("Charlie", encoding="utf-8")

        result = stage._read_artifacts(str(tmp_path))
        assert "Alpha" in result["content"]
        assert "key: value" in result["content"]
        assert "Charlie" in result["content"]
        assert len(result["read"]) == 3

    def test_read_artifacts_reads_all_extensions(self, tmp_path):
        """No extension filter — .vtt, .csv, .docx, etc. are all read."""
        from azext_prototype.stages.design_stage import DesignStage

        stage = DesignStage()
        (tmp_path / "transcript.vtt").write_text("WEBVTT\n00:00 Hello", encoding="utf-8")
        (tmp_path / "notes.rst").write_text("=====\nNotes", encoding="utf-8")
        (tmp_path / "data.csv").write_text("a,b\n1,2", encoding="utf-8")

        result = stage._read_artifacts(str(tmp_path))
        assert "WEBVTT" in result["content"]
        assert "Notes" in result["content"]
        assert "a,b" in result["content"]
        assert len(result["read"]) == 3
        assert result["failed"] == []

    def test_read_artifacts_empty_directory(self, tmp_path):
        from azext_prototype.stages.design_stage import DesignStage

        stage = DesignStage()
        empty_dir = tmp_path / "empty"
        empty_dir.mkdir()

        result = stage._read_artifacts(str(empty_dir))
        assert result["content"] == ""
        assert result["read"] == []
        assert result["failed"] == []

    def test_read_artifacts_nonexistent_path_raises(self, tmp_path):
        from azext_prototype.stages.design_stage import DesignStage

        stage = DesignStage()
        with pytest.raises(CLIError, match="not found"):
            stage._read_artifacts(str(tmp_path / "nonexistent"))

    def test_read_artifacts_nested_directory(self, tmp_path):
        from azext_prototype.stages.design_stage import DesignStage

        stage = DesignStage()
        sub = tmp_path / "sub"
        sub.mkdir()
        (sub / "nested.md").write_text("Nested content", encoding="utf-8")

        result = stage._read_artifacts(str(tmp_path))
        assert "Nested content" in result["content"]
        # Relative path should include subdirectory
        assert any("sub" in r for r in result["read"])

    def test_read_artifacts_binary_image(self, tmp_path):
        """Standalone images are collected in result['images']."""
        from azext_prototype.stages.design_stage import DesignStage

        stage = DesignStage()
        (tmp_path / "notes.md").write_text("# Notes", encoding="utf-8")
        (tmp_path / "arch.png").write_bytes(b"\x89PNG\r\n\x1a\n" + b"\x00" * 50)

        result = stage._read_artifacts(str(tmp_path))
        assert "Notes" in result["content"]
        assert len(result["images"]) == 1
        assert result["images"][0]["mime"] == "image/png"
        assert len(result["read"]) == 2

    def test_read_artifacts_embedded_images(self, tmp_path):
        """Embedded images from documents are collected in result['images']."""
        from azext_prototype.stages.design_stage import DesignStage
        from docx import Document
        from docx.shared import Inches
        from PIL import Image as PILImage
        import io

        stage = DesignStage()
        # Create a DOCX with an embedded image
        img_buf = io.BytesIO()
        PILImage.new("RGB", (10, 10), color="green").save(img_buf, format="PNG")
        img_path = tmp_path / "temp_img.png"
        img_path.write_bytes(img_buf.getvalue())

        doc = Document()
        doc.add_paragraph("With embedded image")
        doc.add_picture(str(img_path), width=Inches(1))
        doc.save(str(tmp_path / "spec.docx"))
        img_path.unlink()  # Remove temp image, only DOCX remains

        result = stage._read_artifacts(str(tmp_path))
        assert "With embedded image" in result["content"]
        assert len(result["images"]) >= 1
        assert any("spec.docx" in img["filename"] for img in result["images"])

    def test_read_artifacts_document_extraction(self, tmp_path):
        """PDF/DOCX text is extracted and included in content."""
        from azext_prototype.stages.design_stage import DesignStage
        from docx import Document

        stage = DesignStage()
        doc = Document()
        doc.add_paragraph("Requirements from Word doc")
        doc.save(str(tmp_path / "req.docx"))

        result = stage._read_artifacts(str(tmp_path))
        assert "Requirements from Word doc" in result["content"]
        assert "req.docx" in result["read"]

    def test_read_artifacts_images_key_present(self, tmp_path):
        """Result dict always has an 'images' key."""
        from azext_prototype.stages.design_stage import DesignStage

        stage = DesignStage()
        (tmp_path / "a.txt").write_text("hello", encoding="utf-8")

        result = stage._read_artifacts(str(tmp_path))
        assert "images" in result
        assert result["images"] == []


class TestDesignStageReadFile:
    """Cover _read_file — now returns ReadResult."""

    def test_read_file_success(self, tmp_path):
        from azext_prototype.stages.design_stage import DesignStage
        from azext_prototype.parsers.binary_reader import FileCategory

        stage = DesignStage()
        f = tmp_path / "test.txt"
        f.write_text("Hello world", encoding="utf-8")
        result = stage._read_file(f)
        assert result.category == FileCategory.TEXT
        assert result.text == "Hello world"
        assert result.error is None

    def test_read_file_unreadable(self, tmp_path):
        from azext_prototype.stages.design_stage import DesignStage

        stage = DesignStage()
        d = tmp_path / "adir"
        d.mkdir()
        result = stage._read_file(d)
        assert result.error is not None

    def test_read_file_image(self, tmp_path):
        from azext_prototype.stages.design_stage import DesignStage
        from azext_prototype.parsers.binary_reader import FileCategory

        stage = DesignStage()
        f = tmp_path / "photo.jpg"
        f.write_bytes(b"\xff\xd8\xff\xe0" + b"\x00" * 50)
        result = stage._read_file(f)
        assert result.category == FileCategory.IMAGE
        assert result.image_data is not None

    def test_read_file_document(self, tmp_path):
        from azext_prototype.stages.design_stage import DesignStage
        from azext_prototype.parsers.binary_reader import FileCategory
        from docx import Document

        stage = DesignStage()
        doc = Document()
        doc.add_paragraph("Test content")
        docx_path = tmp_path / "doc.docx"
        doc.save(str(docx_path))
        result = stage._read_file(docx_path)
        assert result.category == FileCategory.DOCUMENT
        assert "Test content" in result.text


class TestDesignStageLoadSaveState:
    """Cover _load_design_state, _save_design_state."""

    def test_load_design_state_fresh(self, tmp_path):
        from azext_prototype.stages.design_stage import DesignStage

        stage = DesignStage()
        state = stage._load_design_state(str(tmp_path), reset=False)
        assert state["iteration"] == 0

    def test_load_design_state_reset(self, tmp_path):
        from azext_prototype.stages.design_stage import DesignStage

        stage = DesignStage()
        state_path = tmp_path / ".prototype" / "state" / "design.json"
        state_path.parent.mkdir(parents=True)
        state_path.write_text(json.dumps({"iteration": 5}), encoding="utf-8")

        state = stage._load_design_state(str(tmp_path), reset=True)
        assert state["iteration"] == 0

    def test_load_design_state_existing(self, tmp_path):
        from azext_prototype.stages.design_stage import DesignStage

        stage = DesignStage()
        state_path = tmp_path / ".prototype" / "state" / "design.json"
        state_path.parent.mkdir(parents=True)
        existing = {"iteration": 3, "architecture": "## Arch v3"}
        state_path.write_text(json.dumps(existing), encoding="utf-8")

        state = stage._load_design_state(str(tmp_path), reset=False)
        assert state["iteration"] == 3

    def test_load_design_state_corrupt_json(self, tmp_path):
        from azext_prototype.stages.design_stage import DesignStage

        stage = DesignStage()
        state_path = tmp_path / ".prototype" / "state" / "design.json"
        state_path.parent.mkdir(parents=True)
        state_path.write_text("not valid json {{{", encoding="utf-8")

        state = stage._load_design_state(str(tmp_path), reset=False)
        assert state["iteration"] == 0

    def test_save_design_state(self, tmp_path):
        from azext_prototype.stages.design_stage import DesignStage

        stage = DesignStage()
        state = {"iteration": 2, "architecture": "## Arch", "decisions": []}

        stage._save_design_state(str(tmp_path), state)

        state_path = tmp_path / ".prototype" / "state" / "design.json"
        assert state_path.exists()
        loaded = json.loads(state_path.read_text(encoding="utf-8"))
        assert loaded["iteration"] == 2

    def test_save_design_state_creates_directories(self, tmp_path):
        from azext_prototype.stages.design_stage import DesignStage

        stage = DesignStage()
        new_dir = tmp_path / "brand_new"
        new_dir.mkdir()

        stage._save_design_state(str(new_dir), {"iteration": 1})
        assert (new_dir / ".prototype" / "state" / "design.json").exists()


class TestDesignStageWriteArchDocs:
    """Cover _write_architecture_docs."""

    def test_write_architecture_docs(self, tmp_path):
        from azext_prototype.stages.design_stage import DesignStage

        stage = DesignStage()
        stage._write_architecture_docs(str(tmp_path), "# Architecture\nGreat design")

        arch = tmp_path / "concept" / "docs" / "ARCHITECTURE.md"
        assert arch.exists()
        assert "Great design" in arch.read_text(encoding="utf-8")


# ======================================================================
# DeployStage — targeted coverage
# ======================================================================


class TestDeploymentOutputCapture:
    """Cover DeploymentOutputCapture from deploy_helpers."""

    def test_capture_terraform_outputs(self, tmp_path):
        from azext_prototype.stages.deploy_helpers import DeploymentOutputCapture

        capture = DeploymentOutputCapture(str(tmp_path))

        with patch("azext_prototype.stages.deploy_helpers.subprocess.run") as mock_run:
            mock_run.return_value = MagicMock(
                returncode=0,
                stdout='{"rg_name": {"value": "my-rg", "type": "string"}}',
            )
            result = capture.capture_terraform(tmp_path / "concept" / "infra" / "terraform")
            assert result == {"rg_name": "my-rg"}

    def test_capture_bicep_outputs(self, tmp_path):
        from azext_prototype.stages.deploy_helpers import DeploymentOutputCapture

        capture = DeploymentOutputCapture(str(tmp_path))

        deployment_output = '{"properties":{"outputs":{"rg":{"value":"my-rg"}}}}'
        result = capture.capture_bicep(deployment_output)
        assert result == {"rg": "my-rg"}

    def test_capture_bicep_empty_output_returns_empty(self, tmp_path):
        from azext_prototype.stages.deploy_helpers import DeploymentOutputCapture

        capture = DeploymentOutputCapture(str(tmp_path))

        result = capture.capture_bicep("")
        assert result == {}


class TestDeployHelpersBicep:
    """Cover deploy_bicep and deploy_terraform from deploy_helpers."""

    @patch("azext_prototype.stages.deploy_helpers.subprocess.run")
    def test_deploy_bicep_success(self, mock_run, tmp_path):
        from azext_prototype.stages.deploy_helpers import deploy_bicep

        (tmp_path / "main.bicep").write_text("resource x {}", encoding="utf-8")
        mock_run.return_value = MagicMock(returncode=0, stdout='{"outputs":{}}', stderr="")

        result = deploy_bicep(tmp_path, "sub-123", "my-rg")
        assert result["status"] == "deployed"
        assert result["scope"] == "resourceGroup"

    @patch("azext_prototype.stages.deploy_helpers.subprocess.run")
    def test_deploy_terraform_success(self, mock_run, tmp_path):
        from azext_prototype.stages.deploy_helpers import deploy_terraform

        mock_run.return_value = MagicMock(returncode=0, stdout="", stderr="")
        result = deploy_terraform(tmp_path, "sub-123")
        assert result["status"] == "deployed"
        assert result["tool"] == "terraform"

    @patch("azext_prototype.stages.deploy_helpers.subprocess.run")
    def test_deploy_terraform_failure(self, mock_run, tmp_path):
        from azext_prototype.stages.deploy_helpers import deploy_terraform

        mock_run.side_effect = [
            MagicMock(returncode=0, stdout="", stderr=""),  # init
            MagicMock(returncode=1, stdout="", stderr="plan failed"),  # plan
        ]
        result = deploy_terraform(tmp_path, "sub-123")
        assert result["status"] == "failed"
        assert "plan failed" in result["error"]


class TestWhatIfBicep:
    """Cover whatif_bicep from deploy_helpers."""

    @patch("azext_prototype.stages.deploy_helpers.subprocess.run")
    def test_whatif_subscription_scope(self, mock_run, tmp_path):
        from azext_prototype.stages.deploy_helpers import whatif_bicep

        (tmp_path / "main.bicep").write_text(
            "targetScope = 'subscription'\nresource rg 'Microsoft.Resources/resourceGroups@2023-07-01' = {}",
            encoding="utf-8",
        )

        mock_run.return_value = MagicMock(
            returncode=0,
            stdout="Resource changes: 1 to create",
            stderr="",
        )

        result = whatif_bicep(tmp_path, "sub-123", "")
        assert result["status"] == "previewed"

    @patch("azext_prototype.stages.deploy_helpers.subprocess.run")
    def test_whatif_with_params_file(self, mock_run, tmp_path):
        from azext_prototype.stages.deploy_helpers import whatif_bicep

        (tmp_path / "main.bicep").write_text("resource x {}", encoding="utf-8")
        (tmp_path / "main.parameters.json").write_text('{"parameters":{}}', encoding="utf-8")

        mock_run.return_value = MagicMock(returncode=0, stdout="preview ok", stderr="")

        whatif_bicep(tmp_path, "sub-123", "rg")
        cmd_parts = mock_run.call_args[0][0]
        assert "--parameters" in cmd_parts


class TestGetCurrentSubscription:
    """Cover get_current_subscription from deploy_helpers."""

    @patch("azext_prototype.stages.deploy_helpers.subprocess.run")
    def test_get_current_subscription_success(self, mock_run):
        from azext_prototype.stages.deploy_helpers import get_current_subscription

        mock_run.return_value = MagicMock(returncode=0, stdout="aaaabbbb-1234\n")
        assert get_current_subscription() == "aaaabbbb-1234"

    @patch("azext_prototype.stages.deploy_helpers.subprocess.run", side_effect=FileNotFoundError)
    def test_get_current_subscription_file_not_found(self, mock_run):
        from azext_prototype.stages.deploy_helpers import get_current_subscription

        assert get_current_subscription() == ""


class TestDeployBicepSubscriptionScope:
    """Cover deploy_bicep subscription scope with params."""

    @patch("azext_prototype.stages.deploy_helpers.subprocess.run")
    def test_deploy_bicep_subscription_scope(self, mock_run, tmp_path):
        from azext_prototype.stages.deploy_helpers import deploy_bicep

        (tmp_path / "main.bicep").write_text("targetScope = 'subscription'\n", encoding="utf-8")
        (tmp_path / "main.parameters.json").write_text('{"parameters":{}}', encoding="utf-8")

        mock_run.return_value = MagicMock(returncode=0, stdout='{"properties":{}}', stderr="")

        result = deploy_bicep(tmp_path, "sub-123", "")
        assert result["status"] == "deployed"
        assert result["scope"] == "subscription"

        cmd_parts = mock_run.call_args[0][0]
        assert "--parameters" in cmd_parts


class TestDeployAppStage:
    """Cover deploy_app_stage from deploy_helpers."""

    @patch("azext_prototype.stages.deploy_helpers.subprocess.run")
    def test_deploy_app_stage_deploy_script(self, mock_run, tmp_path):
        from azext_prototype.stages.deploy_helpers import deploy_app_stage

        (tmp_path / "deploy.sh").write_text("#!/bin/bash\necho ok", encoding="utf-8")

        mock_run.return_value = MagicMock(returncode=0, stdout="ok", stderr="")

        result = deploy_app_stage(tmp_path, "sub", "rg")
        assert result["status"] == "deployed"
        assert result["method"] == "deploy_script"

    @patch("azext_prototype.stages.deploy_helpers.subprocess.run")
    def test_deploy_app_stage_sub_apps(self, mock_run, tmp_path):
        from azext_prototype.stages.deploy_helpers import deploy_app_stage

        # No top-level deploy.sh, but sub-app directories
        web = tmp_path / "web"
        web.mkdir()
        (web / "deploy.sh").write_text("#!/bin/bash\necho web", encoding="utf-8")

        api = tmp_path / "api"
        api.mkdir()
        (api / "deploy.sh").write_text("#!/bin/bash\necho api", encoding="utf-8")

        mock_run.return_value = MagicMock(returncode=0, stdout="ok", stderr="")

        result = deploy_app_stage(tmp_path, "sub", "rg")
        assert result["status"] == "deployed"
        assert "web" in result["apps"]
        assert "api" in result["apps"]

    @patch("azext_prototype.stages.deploy_helpers.subprocess.run")
    def test_deploy_app_stage_sub_app_failure(self, mock_run, tmp_path):
        """Failed sub-app is logged but doesn't stop others."""
        from azext_prototype.stages.deploy_helpers import deploy_app_stage

        api = tmp_path / "api"
        api.mkdir()
        (api / "deploy.sh").write_text("#!/bin/bash\nexit 1", encoding="utf-8")

        web = tmp_path / "web"
        web.mkdir()
        (web / "deploy.sh").write_text("#!/bin/bash\necho ok", encoding="utf-8")

        mock_run.side_effect = [
            MagicMock(returncode=1, stdout="", stderr="api failed"),
            MagicMock(returncode=0, stdout="ok", stderr=""),
        ]

        result = deploy_app_stage(tmp_path, "sub", "rg")
        assert result["status"] == "deployed"
        assert "web" in result["apps"]
